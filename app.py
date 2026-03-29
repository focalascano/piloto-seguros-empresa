import streamlit as st
import os
from PIL import Image
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

current_dir = os.path.dirname(os.path.abspath(__file__))
logo_path = os.path.join(current_dir, "logotrenes.png")

logo_exists = os.path.exists(logo_path)
logo_image = None

if logo_exists:
    try:
        logo_image = Image.open(logo_path)
    except Exception as e:
        # Si la imagen está corrupta o hay error, no hacemos nada aquí
        pass

# --- 3. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="""Determinador de seguros a proveedores: Generador automático de anexo contractual""",
    page_icon="🛡️", 
    layout="centered"
)

col1, col2 = st.columns([1, 4])

with col1:
    if logo_exists and logo_image:
        st.image(logo_image, width=140)
    else:
        st.title("🛡️")

with col2:
    st.markdown("<h1 style='margin-top: 0;'>Determinador de seguros a proveedores: generador automático de anexo contraactual</h1>", unsafe_allow_html=True)
    

# --- BIBLIOTECA DE CLÁUSULAS (TEXTOS COMPLETOS) ---
TEXTOS_LEGALES = {
    "GENERAL": {
        "encabezado": "La Contratista deberá acreditar ante La SOFSA, con una antelación mínima de CINCO (5) días corridos al inicio de los trabajos y/o servicios, la contratación y vigencia de los seguros que resulten aplicables en función de la naturaleza y riesgos de la prestación, debiendo exigir el cumplimiento de esta obligación a los Subcontratistas que eventualmente participen en la ejecución de sus obligaciones contractuales, cuando la contratación así lo permita:",
        "requisitos": """Otros Seguros: SOFSA se reserva el derecho de exigir otros seguros que, en virtud de la contratación pudiesen ser requeridos.
**Requisitos de los Seguros**  
Las aseguradoras contratadas deberán cumplir con las siguientes condiciones:  
• Ser una aseguradora habilitada por la Superintendencia de Seguros de la Nación.
• Estar calificada por alguna de las Calificadoras de Riesgo autorizadas por la Comisión Nacional de Valores (CNV).  
Se tomará como válida la calificación del año en que se adjudique la contratación y/o la calificación del año inmediato anterior a la adjudicación.  
La Contratista deberá presentar a la Licitante la calificación de riesgos de la Aseguradora.
La Contratista deberá mantener y pagar el premio correspondiente a las pólizas. Los comprobantes de pago de las mismas deberán ser presentados a la Licitante de manera mensual y consecutiva.""",
        "vigencia": """Vigencia de los Seguros  
        Los seguros deberán mantenerse vigentes desde el inicio de cualquier actividad vinculada a la contratación, incluyendo tareas previas, y hasta la extinción total de las obligaciones contractuales de la Contratista, comprendiendo la recepción provisoria, el período de garantía y toda intervención posterior vinculada al contrato.  
        """,
        "responsabilidad": """
        Incumplimientos en la Presentación de los Seguros  
        Si la Contratista no presentase los seguros que correspondan de acuerdo con la naturaleza de la actividad, los trabajos y/o los servicios a ejecutar, o no cumpliera con alguno de los requisitos establecidos en el presente Anexo, no podrá iniciar ni continuar las tareas hasta tanto regularice dicha situación, siendo de su exclusiva responsabilidad las consecuencias que ello genere, sin que ello otorgue derecho a reclamo alguno contra la SOFSA.
Criterio de interpretación y aplicación: Ante cualquier duda razonable respecto de la aplicabilidad, alcance o suficiencia de los seguros exigidos en el presente Anexo, SOFSA tendrá la facultad de definir el seguro que resulte exigible, en función de la naturaleza de la prestación y de los riesgos involucrados.  

Responsabilidad  

La contratación de seguros por parte de la Contratista no limita ni reduce en modo alguno su responsabilidad contractual ni legal, siendo ésta responsable directa por todos los daños y obligaciones derivados de la ejecución del contrato.  
En consecuencia, la Contratista asumirá a su exclusivo cargo las franquicias, descubiertos, diferencias de suma asegurada y todo daño o reclamo que no resulte cubierto por las pólizas contratadas.  
La Contratista mantendrá indemne a SOFSA, ADIFSA, FASE - en proceso de transformación a Sociedad Anónima Unipersonal (SAU)-, Secretaria de Transporte de la Nación, y/o al Estado Nacional, así como a sus accionistas, directores, empleados y funcionarios, frente a cualquier reclamo, suma, daño o gasto que deban afrontar con motivo de la ejecución contractual y/o del incumplimiento del régimen de seguros."""
    },
    "RC": """
    Seguros a presentar por la Contratista cuando, por la naturaleza de la actividad a desarrollar, exista riesgo de ocasionar daños a personas y/o a bienes de terceros:  
    
    Seguro de Responsabilidad Civil Comprensiva:   
    
    La Contratista deberá contratar y mantener vigente, por su exclusiva cuenta y cargo, un seguro de Responsabilidad Civil Comprensiva que deberá cubrir los daños a personas y/o bienes de terceros derivados directa o indirectamente de la ejecución de los trabajos y/o servicios contratados. En caso de insuficiencia o falta de cobertura, los daños deberán ser asumidos íntegramente por la Contratista. Ante el pago de un siniestro, la suma asegurada deberá ser repuesta dentro de los DIEZ (10) días de producido el mismo.  
    
    Coberturas adicionales (condicionales):  
    La póliza deberá incluir, cuando el riesgo asociado a la actividad lo requiera, los adicionales correspondientes a uso de grúas, izaje, andamios, trabajos de soldadura u oxicorte, carga y descarga, maquinaria, transporte de bienes, contaminación súbita y accidental, suministro de alimentos, uso de armas de fuego, uso de vehículos propios o no propios en exceso de su póliza específica y personas físicas bajo contrato. 
    Previo al inicio de las tareas, la Contratista deberá presentar certificado de cobertura y libre deuda emitido por la aseguradora.   
    Clausulas obligatorias: 
    Asegurado Adicional: Serán considerados asegurados y/o asegurados adicionales el titular de la póliza y/o la empresa CUIT 30-71068177-1 y/o ADMINISTRACION DE INFRAESTRUCTURAS FERROVIARIAS SOCIEDAD ANONIMA (ADIFSA) CUIT 30- 71069599-3, y/o FERROCARRILES ARGENTINOS SOCIEDAD DEL ESTADO (FASE) - en proceso de transformación a Sociedad Anónima Unipersonal (SAU) - CUIT 30-71525570-3, y/o a SECRETARIA DE TRANSPORTE DE LA NACIÓN CUIT 30-71512720-9, y/o MINISTERIO DE ECONOMÍA CUIT 30-54667611-7, y/o al ESTADO NACIONAL, quienes serán coasegurados y/o asegurados adicionales a los efectos de la cobertura de la póliza, así como sus accionistas, directores, empleados y funcionarios.   
    Responsabilidad Civil Cruzada: Todos los sujetos mencionados precedentemente serán considerados terceros entre sí.   
    Cláusula de No Repetición: La Aseguradora renunciará expresamente a todo derecho de subrogación o repetición contra los sujetos mencionados precedentemente, manteniendo indemne a la empresa frente a reclamos de terceros cubiertos por la póliza.   
    Notificación previa: La póliza no será anulada sin previo aviso por escrito a la OPERADORA FERROVIARIA SOCIEDAD ANONIMA, con domicilio en la Avda. Ramos Mejía Nº 1302, piso 4to. de la Ciudad Autónoma de Buenos Aires, con un plazo mínimo de 15 días corridos de anticipación.""",
    "ART": """
    Seguros a presentar por la Contratista para el personal que se encuentre en relación de dependencia, y que deba ingresar a predio de SOFSA en virtud de la presente contratación:  
    Seguro de Riesgos del Trabajo:  
    La Contratista deberá contratar y mantener vigente, por su exclusiva cuenta y cargo, un seguro que cubra los riesgos del trabajo de acuerdo con la Ley Nº 24.557 de Riesgos del Trabajo, sus reformas y decretos reglamentarios.  
    
    Previo al inicio de las tareas, la Contratista deberá presentar certificado de cobertura emitido por la ART, incluyendo la nómina del personal afectado. 
    La póliza deberá incluir la siguiente cláusula:  
    Cláusula de No Repetición: La Aseguradora de Riesgos del Trabajo debe renunciar en forma expresa a sus derechos de subrogación y/o a reclamar o iniciar toda acción de repetición o de regreso contra SOFSA, y/o FASE - en proceso de transformación a Sociedad Anónima Unipersonal (SAU) - y/o ADIFSA y/o SECRETARIA DE TRANSPORTE DE LA NACIÓN, y/o MINISTERIO DE ECONOMÍA, y/o ESTADO NACIONAL así como sus accionistas, directores, empleados y funcionarios, con motivo de las prestaciones a las que se vea obligada a otorgar o abonar al personal dependiente o ex dependiente de la Contratista, amparados por la cobertura del contrato de afiliación, por accidente de trabajo o enfermedades profesionales ocurridos o contraídos por el hecho o en ocasión del trabajo o en el trayecto entre el domicilio del trabajador y el lugar de trabajo.  
""",
    "VO": """
    Seguro Colectivo de Vida Obligatorio:  
    La Contratista deberá contratar y mantener vigente, por su exclusiva cuenta y cargo, un seguro colectivo de vida obligatorio para cubrir la totalidad del personal afectado al trabajo y/o servicio contratado, según lo previsto en el Decreto Nº 1567/74. 
    Previo al inicio de las tareas, deberá presentarse certificado de cobertura emitido por la aseguradora, con indicación de la nómina del personal cubierto.
""",
    "AP": """
    Seguros a presentar por la Contratista para el personal contratado que NO se encuentre en relación de dependencia, y que deba ingresar a predio de SOFSA en virtud de la presente contratación:  
    Seguro de Accidentes Personales:  
    La contratista deberá ontratar y mantener vigente, por su exclusiva cuenta y cargo, un seguro que cubra los accidentes que pudiera sufrir el personal de la Contratista, afectado a los trabajos y/o servicios y que no se encuentre en relación de dependencia con ésta, cuando la modalidad contractual así lo permita.  
    La cobertura mínima por persona deberá contemplar:   
    Muerte e incapacidad permanente (total o parcial): USD 20.000 o su equivalente en moneda local.   
    Gastos médicos asistenciales: USD 2.000 o su equivalente en moneda local.   
    La póliza deberá designar a SOFSA como beneficiaria en primer término, exclusivamente a los efectos de garantizar su indemnidad frente a eventuales obligaciones legales derivadas del siniestro. 
    Previo al inicio de las tareas, la Contratista deberá presentar certificado de cobertura y libre deuda emitido por la aseguradora. 
    La póliza deberá incluir las siguientes cláusulas:   
    Notificación previa: La póliza no será anulada sin previo aviso por escrito a la OPERADORA FERROVIARIA SOCIEDAD ANONIMA, con un plazo mínimo de 15 días corridos de anticipación.""",
    "CAUCION": """Caución a presentar por la contratista para aquellos trabajos y/o servicios que requieran el retiro, traslado, transporte, guarda, custodia, reparación, mantenimiento, certificación, service, calibración y/o cualquier tipo de intervención técnica sobre Bienes y/o Equipos, en funcionamiento o no, propiedad de SOFSA y/o bajo su responsabilidad, cuando dichos bienes deban permanecer o ser trasladados fuera de las instalaciones de SOFSA:  
    Caución de Tenencia de Bienes:   
    La Contratista deberá contratar una Póliza de Caución de Tenencia de Bienes, destinada a garantizar el retiro, transporte, traslado, tenencia, guarda, custodia, correcta conservación y posterior devolución de los Bienes y/o Equipos propiedad de SOFSA y/o bajo su responsabilidad.
    La cobertura deberá mantenerse vigente desde el momento del retiro de los bienes de las instalaciones de SOFSA o del lugar que ésta determine, durante todo el período de traslado, permanencia, manipulación, intervención técnica y custodia en instalaciones de la Contratista o de terceros, hasta su efectiva devolución y recepción conforme por parte de SOFSA. 
    La suma asegurada deberá ser equivalente a la suma establecida en el Pliego de Especificaciones Técnicas. 
    La póliza deberá incluir la siguiente cláusula: 
    Los actos, declaraciones, acciones u omisiones de la Contratista (Tomador), incluida la falta de pago del premio, no afectarán de modo alguno los derechos de SOFSA (Asegurada) frente al Asegurador, quien mantendrá íntegramente su obligación de responder en los términos de la póliza.
    """,
    "TRCYM": """
    Seguro a presentar por la Contratista para aquellos trabajos y/o servicios que requieran la realización de una obra y/o montaje:  Seguros Todo riesgo Construcción y/o Montaje:  
    Cuando la contratación implique la ejecución de obras y/o trabajos de montaje, la Contratista deberá contratar y mantener vigente, por su exclusiva cuenta y cargo, un Seguro Todo Riesgo Construcción y/o Montaje que ampare la totalidad de la obra y/o montaje, incluyendo trabajos temporarios, materiales, equipos, instalaciones, obradores, maquinarias y bienes existentes y/o adyacentes afectados a la prestación. 
    La cobertura deberá mantenerse vigente durante todo el período de ejecución, incluyendo los períodos de almacenaje, construcción y/o montaje, pruebas y mantenimiento, y hasta la recepción definitiva de la obra, debiendo actualizarse progresivamente la suma asegurada de modo tal que refleje en todo momento el valor total certificado. 
    La póliza deberá ser contratada a nombre conjunto de la Contratista y de SOFSA, y extenderse, cuando corresponda, a subcontratistas y/o proveedores que intervengan en la ejecución.
    La Contratista deberá presentar a SOFSA certificado de cobertura y libre deuda emitido por la aseguradora.  
    Clausulas obligatorias:
    Asegurado Adicional: Serán considerados asegurados y/o asegurados adicionales el titular de la póliza y/o la La OPERADORA FERROVIARIA SOCIEDAD ANONIMA (SOFSA)  CUIT 30-71068177-1 y/o ADMINISTRACION DE INFRAESTRUCTURAS FERROVIARIAS SOCIEDAD ANONIMA (ADIFSA) CUIT 30- 71069599-3, y/o FERROCARRILES ARGENTINOS SOCIEDAD DEL ESTADO (FASE) - en proceso de transformación a Sociedad Anónima Unipersonal (SAU) - CUIT 30-71525570-3, y/o a SECRETARIA DE TRANSPORTE DE LA NACIÓN CUIT 30-71512720-9, y/o MINISTERIO DE ECONOMÍA CUIT 30-54667611-7, y/o al ESTADO NACIONAL, quienes serán coasegurados y/o asegurados adicionales a los efectos de la cobertura de la póliza, así como sus accionistas, directores, empleados y funcionarios. 
    Responsabilidad Civil Cruzada: Todos los sujetos mencionados precedentemente serán considerados terceros entre sí.
    Cláusula de No Repetición: La Aseguradora renunciará expresamente a todo derecho de subrogación o repetición contra los sujetos mencionados precedentemente, manteniendo indemne a SOFSA frente a reclamos de terceros cubiertos por la póliza.
    Notificación previa: La póliza no será anulada sin previo aviso por escrito a la OPERADORA FERROVIARIA SOCIEDAD ANONIMA, con un plazo mínimo de 15 días corridos de anticipación.
    """,
    "AUTO": """
    Seguros a presentar por la Contratista para los vehículos a ser utilizados en virtud de la presente contratación y/o que ingresen a predio de SOFSA:  Seguro Automotor Obligatorio: 
    La Contratista deberá contratar y mantener vigente, por su exclusiva cuenta y cargo, un seguro Automotor para los vehículos a ser utilizados en virtud de la presente contratación, los cuales deberán contar, como mínimo, con la cobertura de Responsabilidad Civil - Seguro Voluntario, por la suma establecida por la Superintendencia de Seguros de la Nación.
    La Contratista deberá presentar a SOFSA un certificado de cobertura y libre deuda emitido por la Aseguradora.
    La póliza deberá incluir las siguientes cláusulas:  
    Cláusula de No Repetición: La Aseguradora renunciará expresamente a todo derecho de subrogación o repetición contra SOFSA, y/o FASE - en proceso de transformación a Sociedad Anónima Unipersonal (SAU) - y/o ADIFSA y/o SECRETARIA DE TRANSPORTE DE LA NACIÓN, y/o MINISTERIO DE ECONOMÍA, y/o ESTADO NACIONAL, así como sus accionistas, directores, empleados y funcionarios, con motivo de las sumas que se vea obligada a abonar por los riesgos amparados en la cobertura de la póliza.
    Notificación previa: La póliza no será anulada sin previo aviso por escrito a la OPERADORA FERROVIARIA SOCIEDAD ANONIMA, con domicilio en la Avda. Ramos Mejía Nº 1302, piso 4to. de la Ciudad Autónoma de Buenos Aires, con un plazo mínimo de 15 días corridos de anticipación.
"""
}

def generar_anexo_completo(seguros_activos, nivel):
    doc = Document()
    
    # --- CONFIGURACIÓN DE ESTILO GLOBAL (Calibri) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    def agregar_parrafo_formateado(texto, negrita=False, es_titulo=False):
        p = doc.add_paragraph()
        # Alineación justificada
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = p.add_run(texto)
        run.bold = negrita
        if es_titulo:
            run.font.size = Pt(14)
        return p

    # Título principal
    agregar_parrafo_formateado('ANEXO DE SEGUROS', negrita=True, es_titulo=True)
    
    # Encabezado General
    agregar_parrafo_formateado(TEXTOS_LEGALES["GENERAL"]["encabezado"])

    # Sección de Seguros
    if not seguros_activos:
        agregar_parrafo_formateado("No se han determinado seguros específicos adicionales bajo el nivel de Riesgo Nulo.")
    else:
        for s in seguros_activos:
            # Cláusula (Justificada)
            agregar_parrafo_formateado(s['clausula'])
            
            # Suma Asegurada
            if 'suma' in s:
                p_suma = agregar_parrafo_formateado(f"SUMA ASEGURADA MÍNIMA REQUERIDA: {s['suma']}", negrita=False)
            
            agregar_parrafo_formateado("_" * 30)

    # Requisitos, Vigencia y Responsabilidad
    agregar_parrafo_formateado(TEXTOS_LEGALES["GENERAL"]["requisitos"])
    agregar_parrafo_formateado(TEXTOS_LEGALES["GENERAL"]["vigencia"])
    agregar_parrafo_formateado(TEXTOS_LEGALES["GENERAL"]["responsabilidad"])

    # --- CORRECCIÓN PARA FORZAR CALIBRI EN TODO EL DOC ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- INFORMACIÓN DE LA VERSIÓN Y AUTORES ---
st.write("""
Responda el siguiente cuestionario para determinar los seguros aplicables. 
Al final se podrá descargar un archivo Word con el Anexo de seguros que corresponda contratar.
""")

st.caption("""
**Versión 1.0 – 2026** **Autores:** Diego Martín Morris, Ignacio Khoury.  
**Colaboradores:** Claudia Grahl, Facundo Gonzalez, Gonzalo Dince, Analia Zalazar, Silvina Coronel.   
**Implementación informatica del moodelo:** Fernanda Lascano  
**Fecha:** 2026
""")

st.markdown("---")
# Cuestionario Sí/No
opciones = ["No", "Sí"]

r1 = st.radio("¿Para realizar la actividad personal del proveedor ingresará a predios o instalaciones de la empresa?", opciones, index=0)
r2 = st.radio("¿La actividad consiste exclusivamente en tareas administrativas o profesionales de oficina, sin intervención técnica ni operativa?", opciones, index=0)
r3 = st.radio("¿La actividad requiere uso o ingreso de vehículos del proveedor a predios o instalaciones de la empresa?", opciones, index=0)
r4 = st.radio("¿El proveedor transportará o tendrá en sus instalaciones mercadería, bienes o equipos de la empresa?", opciones, index=0)
r5 = st.radio("¿El trabajo se realizará en estaciones, andenes, vías, talleres ferroviarios o sectores con circulación de trenes o pasajeros?", opciones, index=0)
r6 = st.radio("""¿La actividad corresponde a un trabajo menor de mantenimiento simple en la empresa? Para ser considerado trabajo menor, debe cumplir todas estas condiciones:
• duración corta (menor a 1 mes de trabajo)  
• uso herramientas manuales simples  
• sin trabajo en altura, ni andamios  
• sin maquinaria  
• sin intervención en infraestructura  
• sin afectar circulación ferroviaria o de pasajeros  
       **Ejemplos: (pintura interior de oficina, reparación menor de mobiliario, cerrajería, etc)""", opciones, index=0)
r7 = st.radio("¿La actividad requiere uso de equipos, maquinaria o de herramientas complejas en la empresa? Ejemplos: herramientas de corte y/o herramienta de calor y/o herramienta a explosión, equipos técnicos, maquinarias", opciones, index=0)
r8 = st.radio("""¿La actividad incluye alguna de las siguientes tareas  
• trabajos en altura  
• soldadura u oxicorte  
• izaje de cargas  
• intervención eléctrica  
• uso de maquinaria pesada  
• uso de armas de fuego  
• suministro de alimentos?""", opciones, index=0)
r9 = st.radio("""¿La actividad implica construir, instalar o montar una obra, sistema o equipos nuevo? 
Incluye:  
• obras civiles,  
• refacciones estructurales  
• instalación de equipos (montaje o desmontaje)  
• montaje de sistema electrico o mecánico  
No incluye:  
• mantenimiento simple  
• refacciones menores  
• tareas de servicio""", opciones, index=0)

# Mapeo a booleanos para lógica interna
p1 = (r1 == "Sí")
p2 = (r2 == "Sí")
p3 = (r3 == "Sí")
p4 = (r4 == "Sí")
p5 = (r5 == "Sí")
p6 = (r6 == "Sí")
p7 = (r7 == "Sí")
p8 = (r8 == "Sí")
p9 = (r9 == "Sí")

# Validaciones
errores = []
if not p1 and (p3 or p4 or p5 or p6 or p7 or p8 or p9):
    errores.append("⚠️ Bloqueo: Toda condición operativa requiere el ingreso de personal (P1 = SÍ).")
if p2 and (p4 or p5 or p6 or p7 or p8 or p9):
    errores.append("⚠️ Bloqueo: Tareas administrativas incompatibles con riesgos operativos.")
if p6 and (p4 or p5 or p7 or p8 or p9):
    errores.append("⚠️ Bloqueo: Trabajo menor incompatible con riesgos altos o maquinaria compleja.")

if st.button("Generar Documento Final"):
    if errores:
        for err in errores: 
            st.error(err)
    else:
        # Nivel de Riesgo
        if not p1: 
            nivel = "Nulo"
        elif p9 or p8 or p5 or p4: 
            nivel = "Alto"
        elif p1 and p7: 
            nivel = "Medio"
        else: 
            nivel = "Bajo"

        seguros_para_word = []

        if p1:
            # Bloque base para p1
            seguros_para_word.append({
                'clausula': f"{TEXTOS_LEGALES['ART']}\n\n{TEXTOS_LEGALES['VO']}\n\n{TEXTOS_LEGALES['AP']}", 
            })
            
            # Condicionales dependientes de p1
            if p5 or p7 or r8 == "Sí" or p9:
                suma_rc = "USD 100.000 (o eq. local)" if nivel == "Alto" else "USD 50.000 (o eq. local)"
                seguros_para_word.append({'clausula': TEXTOS_LEGALES["RC"], 'suma': suma_rc})

        # Bloques independientes
        if p4:
            seguros_para_word.append({'clausula': TEXTOS_LEGALES["CAUCION"]})

        if p9:
            seguros_para_word.append({'clausula': TEXTOS_LEGALES["TRCYM"]})

        if p3:
            seguros_para_word.append({'clausula': TEXTOS_LEGALES["AUTO"]})

        # --- Lógica de Colores para el Cartel ---
        mensaje_nivel = f"Nivel de Riesgo Determinado: {nivel}"
        
        if nivel == "Nulo":
            st.info(mensaje_nivel)  # Azul gris
        elif nivel == "Bajo":
            st.success(mensaje_nivel)     # Verde
        elif nivel == "Medio":
            st.warning(mensaje_nivel)  # Naranja/Amarillo
        elif nivel == "Alto":
            st.error(mensaje_nivel)    # Rojo

        docx_data = generar_anexo_completo(seguros_para_word, nivel)
        st.download_button("📥 Descargar Anexo Word", docx_data, f"Anexo_Seguros_{nivel}.docx")
